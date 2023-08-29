import os
import re
import telebot
import json
from telebot import types
from num2words import num2words
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from month import *
from msgs import *
from t0ken import *

bot = telebot.TeleBot(TOKEN)

print("run in progress...")
numero = 0
placeholders = {}
subs = []
zap = ''
call_gph = ('nomer', 'date', 'date_ok', 'stoim', 'fio', 'pasp', 'kod', 'dr', 'vyd', 'kem', 'adr', 'rs', 'bik',
            'mail', 'tel')
call_sz = ('nomer', 'date', 'date_ok', 'fio', 'pasp', 'kod', 'dr', 'vyd', 'kem', 'adr', 'rs', 'bik', 'inn', 'raboty',
           'diplom',
           'stoim', 'genzak', 'data_dog', 'nom_dog', 'ikz', 'spr_god', 'spr_data', 'spr_nomer', 'mail_zak', 'tel_zak',
           'mail', 'tel')


@bot.message_handler(commands=['start'])
def start(message):
    try:
        markup = types.InlineKeyboardMarkup(row_width=1)
        item1 = types.InlineKeyboardButton('Договор ГПХ', callback_data='gph')
        item2 = types.InlineKeyboardButton('Договор подряда', callback_data='sz')
        markup.add(item1, item2)
        bot.send_message(message.chat.id, text="Какой документ заполняем?",
                         parse_mode='Markdown', reply_markup=markup)
    except Exception as e:
        print(e)


@bot.message_handler(content_types=['text'])
def answer(message):
    markup = types.InlineKeyboardMarkup(row_width=1)
    if zap == "ГПХ":
        placeholders[subs[numero]] = message.text
        if 0 < numero < len(subs) - 1:
            item1 = types.InlineKeyboardButton('Назад', callback_data=call_gph[numero - 1])
            item2 = types.InlineKeyboardButton('Далее', callback_data=call_gph[numero + 1])
            markup.add(item2, item1)
            bot.reply_to(message, 'Записал, нажмите *"Далее"* для продолжения', parse_mode='Markdown',
                         reply_markup=markup)
        elif numero == 0:
            item1 = types.InlineKeyboardButton('Назад', callback_data='gph')
            item2 = types.InlineKeyboardButton('Далее', callback_data=call_gph[numero + 1])
            markup.add(item2, item1)
            bot.reply_to(message, 'Записал, нажмите *"Далее"* для продолжения', parse_mode='Markdown',
                         reply_markup=markup)
        elif numero == len(subs) - 1:
            item1 = types.InlineKeyboardButton('Назад', callback_data=call_gph[numero - 1])
            item2 = types.InlineKeyboardButton('Готово', callback_data='send')
            markup.add(item2, item1)
            bot.reply_to(message, 'Заполнение закончено. Нажмите *"Готово"* для формирования документа',
                         parse_mode='Markdown', reply_markup=markup)

    elif zap == "СЗ":
        placeholders[subs[numero]] = message.text
        if 0 < numero < len(subs):
            item1 = types.InlineKeyboardButton('Назад', callback_data=call_sz[numero - 1])
            item2 = types.InlineKeyboardButton('Далее', callback_data=call_sz[numero + 1])
            markup.add(item2, item1)
            bot.reply_to(message, 'Записал, нажмите *"Далее"* для продолжения', parse_mode='Markdown',
                         reply_markup=markup)


@bot.callback_query_handler(func=lambda call: call.data == "gph")
def fill(call):
    global zap, subs
    markup = types.InlineKeyboardMarkup(row_width=1)
    item1 = types.InlineKeyboardButton('Приступить к заполнению', callback_data='nomer')
    markup.add(item1)
    bot.send_message(call.message.chat.id, text="Чтобы приступить к заполнению договора ГПХ нажмите на кнопку ниже",
                     reply_markup=markup)

    subs = ["{{номер}}", "{{дата}}", "{{ок_дата}}", "{{стоимость}}", "{{фио}}", "{{серияномер}}", "{{код}}", "{{др}}",
            "{{дата_выд}}", "{{кем}}", "{{адрес}}",
            "{{р/с}}", "{{бик}}", "{{@исполн}}", "{{№исполн}}"]
    zap = 'ГПХ'


@bot.callback_query_handler(func=lambda call: call.data == "sz")
def fill(call):
    global zap, subs
    markup = types.InlineKeyboardMarkup(row_width=1)
    item1 = types.InlineKeyboardButton('Приступить к заполнению', callback_data='nomer')
    markup.add(item1)
    bot.send_message(call.message.chat.id, text="Чтобы приступить к заполнению договора СЗ нажмите на кнопку ниже",
                     reply_markup=markup)

    subs = ["{{номер}}", "{{дата}}", "{{ок_дата}}", "{{фио}}", "{{серияномер}}", "{{код}}", "{{дата}}", "{{кем}}",
            "{{адрес}}", "{{работы}}", "{{кем}}", "{{датадог}}", "{{догном}}", "{{икз}}", "{{гензаказ}}", "{{спр202_}}",
            "{{спрдат}}", "{{спрном}}", "{{стоимость}}", "{{диплом}}", "{{№субчик}}", "{{@исполн}}", "{{№исполн}}",
            "{{№зкзчк}}", "{{ @ зкзчк}}", "{{инн}}", "{{р/с}}"]
    zap = 'СЗ'


# (0, 1, 2, 4, 5, )

@bot.callback_query_handler(func=lambda call: call.data == "nomer")
def nom_dog(call):
    global numero
    numero = 0
    bot.send_message(call.message.chat.id, text="Введите номер договора")


@bot.callback_query_handler(func=lambda call: call.data == "date")
def date(call):
    global numero
    numero = 1
    bot.send_message(call.message.chat.id, text='Введите дату заключения договора (в формате\n(в формате дд.мм.гггг)')


@bot.callback_query_handler(func=lambda call: call.data == "date_ok")
def date_ok(call):
    global numero
    numero = 2
    bot.send_message(call.message.chat.id, text='Введите дату окончания действия договора\n(в формате дд.мм.гггг)')


@bot.callback_query_handler(func=lambda call: call.data == "stoim")
def stoim(call):
    global numero
    numero = 3
    bot.send_message(call.message.chat.id, text="Введите стоимость услуг исполнителя\n(в формате 75000.00)")


@bot.callback_query_handler(func=lambda call: call.data == "fio")
def fio(call):
    global numero
    numero = 4
    bot.send_message(call.message.chat.id, text="Введите ФИО исполнителя")


@bot.callback_query_handler(func=lambda call: call.data == "pasp")
def pasp(call):
    global numero
    numero = 5
    bot.send_message(call.message.chat.id, text="Введите серию, номер паспорта")


@bot.callback_query_handler(func=lambda call: call.data == "kod")
def kod(call):
    global numero
    numero = 6
    bot.send_message(call.message.chat.id, text="Введите код подразделения")


@bot.callback_query_handler(func=lambda call: call.data == "dr")
def dr(call):
    global numero
    numero = 7
    bot.send_message(call.message.chat.id, text="Введите дату роджения\n(в формате дд.мм.гггг)")


@bot.callback_query_handler(func=lambda call: call.data == "vyd")
def vyd(call):
    global numero
    numero = 8
    bot.send_message(call.message.chat.id, text="Введите дату выдачи паспорта\n(в формате дд.мм.гггг)")


@bot.callback_query_handler(func=lambda call: call.data == "kem")
def kem(call):
    global numero
    numero = 9
    bot.send_message(call.message.chat.id, text="Введите кем выдан паспорт")


@bot.callback_query_handler(func=lambda call: call.data == "adr")
def adr(call):
    global numero
    numero = 10
    bot.send_message(call.message.chat.id, text="Введите адрес прописки исполнителя")


@bot.callback_query_handler(func=lambda call: call.data == "rs")
def rs(call):
    global numero
    numero = 11
    bot.send_message(call.message.chat.id, text="Введите рассчетный счет")


@bot.callback_query_handler(func=lambda call: call.data == "bik")
def bik(call):
    global numero
    numero = 12
    bot.send_message(call.message.chat.id, text="Введите БИК банка")


@bot.callback_query_handler(func=lambda call: call.data == "mail")
def mail(call):
    global numero
    numero = 13
    bot.send_message(call.message.chat.id, text="Введите эл. почту исполнителя")


@bot.callback_query_handler(func=lambda call: call.data == "tel")
def tel(call):
    global numero
    numero = 14
    bot.send_message(call.message.chat.id, text="Введите номер телефона исполнителя\n(в формате 7 9xx xxx xx xx)")


@bot.callback_query_handler(func=lambda call: call.data == "send")
def docx(callback_query):
    try:
        global placeholders
        print('сенд')
        placeholders["{{коп}}"] = re.findall(r'.*\.(..)', placeholders["{{стоимость}}"])[0]
        placeholders["{{стоимость}}"] = re.findall(r'(.*)\...', placeholders["{{стоимость}}"])[0]
        placeholders["{{нум_ту_вордс}}"] = num2words(int(placeholders["{{стоимость}}"]), lang='ru')
        # placeholders["{{нум_ту_вордс}}"] = 'Ошибка! Неправильный формат'
        secondname, firstname, otch = placeholders["{{фио}}"].split()
        placeholders["{{фио_кор}}"] = firstname[0] + '. ' + otch[0] + '. ' + secondname
        # placeholders["{{фио_кор}}"] = placeholders["{{фио}}"]

        placeholders["{{банк}}"] = json2dict(placeholders["{{бик}}"])

        print(re.findall(r'(..)\..*', str(placeholders["{{дата}}"])))
        placeholders["{{н_д}}"] = re.findall(r'(..)\..*', str(placeholders["{{дата}}"]))[0]
        print(re.findall(r'..\.(..)\.....', str(placeholders["{{дата}}"]))[0])
        placeholders["{{н_м}}"] = months[re.findall(r'..\.(..)\.....', str(placeholders["{{дата}}"]))[0]]
        print(re.findall(r'..\...\.(.*)', str(placeholders["{{дата}}"])))
        placeholders["{{н_г}}"] = re.findall(r'..\...\.(.*)', str(placeholders["{{дата}}"]))[0]

        placeholders["{{ндфл}}"] = str(round(0.13 * int(placeholders["{{стоимость}}"])))
        placeholders["{{ндфл_ту_вордс}}"] = num2words(int(placeholders["{{ндфл}}"]), lang='ru')
        placeholders["{{к_д}}"] = re.findall(r'(..)\..*', str(placeholders["{{ок_дата}}"]))[0]
        placeholders["{{к_м}}"] = months[re.findall(r'..\.(..)\.....', str(placeholders["{{ок_дата}}"]))[0]]
        placeholders["{{к_г}}"] = re.findall(r'..\...\.(.*)', str(placeholders["{{ок_дата}}"]))[0]

        print(placeholders)
        doc = Document('Договор ГПХ.docx')

        for paragraph in doc.paragraphs:
            for placeholder, value in placeholders.items():
                if placeholder in paragraph.text:
                    # Заменяем текст и устанавливаем шрифт
                    if value:
                        paragraph.text = paragraph.text.replace(placeholder, value)
                        run = paragraph.runs[0]
                        font = run.font
                        font.name = "Times New Roman"
                        font.size = Pt(12)
                        new_doc_name = 'Договор ГПХ ' + secondname + '.docx'
                        doc.save(new_doc_name)

        target_text = placeholders["{{фио}}"]
        placeholders = {}
        for paragraph in doc.paragraphs:
            if target_text in paragraph.text:
                run = paragraph.runs[0]  # Берем первую текстовую руну параграфа
                print(run.text)
                if target_text in run.text:
                    # Разделяем текстовую руну на две части
                    parts = [target_text, run.text.split(target_text)[1]]
                    run.clear()
                    # Добавляем текст перед выделенной частью без жирного начертания
                    new_run = paragraph.add_run(target_text)
                    new_run.bold = True  # Устанавливаем жирное начертание для новой руны
                    font = new_run.font
                    font.name = "Times New Roman"
                    font.size = Pt(12)
                    # Создаем новую руну для выделенной части
                    new_run = paragraph.add_run(parts[1])
                    new_run.bold = False  # Устанавливаем жирное начертание для новой руны

                    font = new_run.font
                    font.name = "Times New Roman"
                    font.size = Pt(12)
                    # Добавляем текст после выделенной части без жирного начертания
                    # new_run.add_text(parts[1])
                    doc.save(new_doc_name)
        with open(new_doc_name, 'rb') as docx_file:
            bot.send_document(callback_query.message.chat.id, docx_file)
    except Exception as e:
        print(e)


if __name__ == "__main__":
    bot.polling()
